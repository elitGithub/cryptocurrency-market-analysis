"""
DOCX Report Generator - Creates Word document reports.
"""
import os
import logging
from typing import Dict, List, Optional
import pandas as pd
from .base_generator import ReportGenerator

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DOCXReportGenerator(ReportGenerator):
    """Generates Word document reports using python-docx library."""

    def generate_node(
        self,
        symbol: str,
        signal_data: Dict,
        exchange_results: pd.DataFrame,
        suggestions: List[str],
        df: Optional[pd.DataFrame],
        chart_path: Optional[str],
        date_prefix: str
    ) -> bool:
        """Generate a Word document report."""
        if not DOCX_AVAILABLE:
            logging.error("python-docx not installed. Run: pip install python-docx")
            return False

        logging.info("Generating Word document report...")

        docx_filename = f"{date_prefix}Crypto_Market_Analysis.docx"
        docx_path = os.path.join(self.output_dir, docx_filename)

        try:
            doc = Document()

            # Title
            title = doc.add_heading('Cryptocurrency Market Analysis Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            subtitle = doc.add_paragraph(symbol)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].font.size = Pt(16)
            subtitle.runs[0].font.color.rgb = RGBColor(127, 140, 141)

            # Executive Summary
            doc.add_heading('Executive Summary', 1)

            signal = signal_data.get('signal', 'N/A')
            confidence = signal_data.get('confidence', 'N/A')
            reasoning = signal_data.get('reasoning', 'N/A')
            price = signal_data.get('price', 0.0)
            rsi = signal_data.get('rsi', 50.0)

            # Signal color
            signal_colors = {'BUY': RGBColor(46, 204, 113), 'SELL': RGBColor(231, 76, 60), 'HOLD': RGBColor(243, 156, 18)}
            signal_color = signal_colors.get(signal, RGBColor(128, 128, 128))

            p = doc.add_paragraph()
            p.add_run('Trading Recommendation: ')
            run = p.add_run(signal)
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = signal_color

            doc.add_paragraph(f'Confidence Level: {confidence}').runs[0].bold = True
            doc.add_paragraph(reasoning).runs[0].italic = True
            doc.add_paragraph(f'Current Price: ${price:,.2f}')
            doc.add_paragraph(f'RSI: {rsi:.1f}')

            # Exchange Comparison
            doc.add_heading('Exchange Comparison', 1)
            doc.add_paragraph('Analysis of cryptocurrency exchanges:')

            if not exchange_results.empty:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Light Grid Accent 1'

                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Exchange'
                hdr_cells[1].text = 'Total Pairs'
                hdr_cells[2].text = 'USDT Pairs'
                hdr_cells[3].text = 'Has OHLCV'

                for _, row in exchange_results.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(row.get('name', 'N/A'))
                    row_cells[1].text = str(row.get('total_spot_pairs', 'N/A'))
                    row_cells[2].text = str(row.get('usdt_quoted_pairs', 'N/A'))
                    row_cells[3].text = 'Yes' if row.get('supports_fetchOHLCV', False) else 'No'

            # Technical Analysis
            doc.add_heading('Technical Analysis Details', 1)
            for suggestion in suggestions:
                doc.add_paragraph(suggestion, style='List Bullet')

            # Chart
            if chart_path and os.path.exists(chart_path):
                doc.add_heading('Price Chart', 2)
                doc.add_picture(chart_path, width=Inches(6))

            # Educational Content
            doc.add_heading('Understanding Technical Indicators', 1)
            doc.add_heading('Moving Averages', 2)
            doc.add_paragraph('Smooth price data to identify trends', style='List Bullet')
            doc.add_paragraph('Golden Cross: Short MA > Long MA (bullish)', style='List Bullet')
            doc.add_paragraph('Death Cross: Short MA < Long MA (bearish)', style='List Bullet')

            doc.add_heading('RSI (Relative Strength Index)', 2)
            doc.add_paragraph('Measures momentum (0-100 scale)', style='List Bullet')
            doc.add_paragraph('RSI > 70: Overbought (potential pullback)', style='List Bullet')
            doc.add_paragraph('RSI < 30: Oversold (potential bounce)', style='List Bullet')

            # Disclaimer
            doc.add_heading('Important Disclaimer', 2)
            disclaimer_p = doc.add_paragraph()
            disclaimer_p.add_run('This report is for educational purposes only').bold = True
            doc.add_paragraph('NOT financial advice', style='List Bullet')
            doc.add_paragraph('Cryptocurrency trading has substantial risk', style='List Bullet')
            doc.add_paragraph('Always do your own research (DYOR)', style='List Bullet')
            doc.add_paragraph('Never invest more than you can afford to lose', style='List Bullet')

            doc.save(docx_path)
            logging.info(f"✓ Word document generated: {docx_filename}")
            return True

        except Exception as e:
            logging.error(f"Error during DOCX generation: {e}")
            return False

    def _generate_js_code(
        self,
        symbol: str,
        signal_data: Dict,
        exchange_results: pd.DataFrame,
        suggestions: List[str],
        chart_path: Optional[str],
        output_path: str
    ) -> str:
        """Generate Node.js code for creating the Word document."""

        # Safely prepare data
        signal = signal_data.get('signal', 'N/A')
        confidence = signal_data.get('confidence', 'N/A')
        reasoning = signal_data.get('reasoning', 'N/A').replace('"', '\\"').replace("'", "\\'")
        price = signal_data.get('price', 0.0)
        rsi = signal_data.get('rsi', 50.0)

        # Prepare exchange data
        exchanges_json = json.dumps(
            exchange_results.to_dict('records') if isinstance(exchange_results, pd.DataFrame) else []
        )

        # Prepare suggestions
        suggestions_json = json.dumps(suggestions if suggestions else [])

        # Determine signal color
        signal_colors = {'BUY': '2ECC71', 'SELL': 'E74C3C', 'HOLD': 'F39C12'}
        signal_color = signal_colors.get(signal, '808080')

        # Check if chart exists
        has_chart = chart_path and os.path.exists(chart_path)
        chart_path_json = json.dumps(chart_path) if has_chart else 'null'

        # Generate JavaScript code
        js_code = f"""
const {{ Document, Packer, Paragraph, TextRun, Table, TableRow, TableCell, ImageRun,
        AlignmentType, HeadingLevel, BorderStyle, WidthType, ShadingType, LevelFormat }} = require('docx');
const fs = require('fs');

const symbol = {json.dumps(symbol)};
const signal = {json.dumps(signal)};
const confidence = {json.dumps(confidence)};
const reasoning = {json.dumps(reasoning)};
const price = {price};
const rsi = {rsi};
const signalColor = {json.dumps(signal_color)};
const exchangeData = {exchanges_json};
const suggestions = {suggestions_json};
const chartPath = {chart_path_json};
const outputPath = {json.dumps(output_path)};

const tableBorder = {{ style: BorderStyle.SINGLE, size: 1, color: "CCCCCC" }};

// Create document
const doc = new Document({{
    styles: {{
        default: {{ document: {{ run: {{ font: "Arial", size: 24 }} }} }},
        paragraphStyles: [
            {{ id: "Title", name: "Title", basedOn: "Normal",
               run: {{ size: 56, bold: true, color: "1C2833", font: "Arial" }},
               paragraph: {{ spacing: {{ before: 240, after: 120 }}, alignment: AlignmentType.CENTER }} }},
            {{ id: "Heading1", name: "Heading 1", basedOn: "Normal", quickFormat: true,
               run: {{ size: 32, bold: true, color: "2E4053", font: "Arial" }},
               paragraph: {{ spacing: {{ before: 240, after: 180 }}, outlineLevel: 0 }} }},
            {{ id: "Heading2", name: "Heading 2", basedOn: "Normal", quickFormat: true,
               run: {{ size: 28, bold: true, color: "34495E", font: "Arial" }},
               paragraph: {{ spacing: {{ before: 180, after: 120 }}, outlineLevel: 1 }} }}
        ]
    }},
    numbering: {{
        config: [
            {{ reference: "bullet-list",
               levels: [{{ level: 0, format: LevelFormat.BULLET, text: "•", alignment: AlignmentType.LEFT,
                          style: {{ paragraph: {{ indent: {{ left: 720, hanging: 360 }} }} }} }}] }}
        ]
    }},
    sections: [{{
        properties: {{ page: {{ margin: {{ top: 1440, right: 1440, bottom: 1440, left: 1440 }} }} }},
        children: [
            // Title
            new Paragraph({{
                heading: HeadingLevel.TITLE,
                children: [new TextRun("Cryptocurrency Market Analysis Report")]
            }}),
            new Paragraph({{
                alignment: AlignmentType.CENTER,
                spacing: {{ after: 400 }},
                children: [new TextRun({{ text: symbol, size: 32, color: "7F8C8D" }})]
            }}),
            
            // Executive Summary
            new Paragraph({{
                heading: HeadingLevel.HEADING_1,
                children: [new TextRun("Executive Summary")]
            }}),
            new Paragraph({{
                spacing: {{ after: 200 }},
                children: [
                    new TextRun("Trading Recommendation: "),
                    new TextRun({{ text: signal, bold: true, size: 32, color: signalColor }})
                ]
            }}),
            new Paragraph({{
                spacing: {{ after: 200 }},
                children: [
                    new TextRun("Confidence Level: "),
                    new TextRun({{ text: confidence, bold: true, size: 24 }})
                ]
            }}),
            new Paragraph({{
                spacing: {{ after: 300 }},
                children: [new TextRun({{ text: reasoning, italics: true }})]
            }}),
            new Paragraph({{
                spacing: {{ after: 200 }},
                children: [
                    new TextRun("Current Price: "),
                    new TextRun({{ text: `$${{price.toFixed(2)}}`, bold: true, size: 28 }})
                ]
            }}),
            new Paragraph({{
                spacing: {{ after: 400 }},
                children: [
                    new TextRun("RSI: "),
                    new TextRun({{ text: rsi.toFixed(1), bold: true, size: 28 }})
                ]
            }}),
            
            // Exchange Comparison
            new Paragraph({{
                heading: HeadingLevel.HEADING_1,
                children: [new TextRun("Exchange Comparison")]
            }}),
            new Paragraph({{
                spacing: {{ after: 200 }},
                children: [new TextRun("Analysis of cryptocurrency exchanges:")]
            }}),
            
            // Exchange Table
            new Table({{
                columnWidths: [2340, 2340, 2340, 2340],
                margins: {{ top: 100, bottom: 100, left: 180, right: 180 }},
                rows: [
                    new TableRow({{
                        tableHeader: true,
                        children: [
                            new TableCell({{
                                borders: {{ top: tableBorder, bottom: tableBorder, left: tableBorder, right: tableBorder }},
                                width: {{ size: 2340, type: WidthType.DXA }},
                                shading: {{ fill: "2E4053", type: ShadingType.CLEAR }},
                                children: [new Paragraph({{
                                    alignment: AlignmentType.CENTER,
                                    children: [new TextRun({{ text: "Exchange", bold: true, color: "FFFFFF", size: 22 }})]
                                }})]
                            }}),
                            new TableCell({{
                                borders: {{ top: tableBorder, bottom: tableBorder, left: tableBorder, right: tableBorder }},
                                width: {{ size: 2340, type: WidthType.DXA }},
                                shading: {{ fill: "2E4053", type: ShadingType.CLEAR }},
                                children: [new Paragraph({{
                                    alignment: AlignmentType.CENTER,
                                    children: [new TextRun({{ text: "Total Pairs", bold: true, color: "FFFFFF", size: 22 }})]
                                }})]
                            }}),
                            new TableCell({{
                                borders: {{ top: tableBorder, bottom: tableBorder, left: tableBorder, right: tableBorder }},
                                width: {{ size: 2340, type: WidthType.DXA }},
                                shading: {{ fill: "2E4053", type: ShadingType.CLEAR }},
                                children: [new Paragraph({{
                                    alignment: AlignmentType.CENTER,
                                    children: [new TextRun({{ text: "USDT Pairs", bold: true, color: "FFFFFF", size: 22 }})]
                                }})]
                            }}),
                            new TableCell({{
                                borders: {{ top: tableBorder, bottom: tableBorder, left: tableBorder, right: tableBorder }},
                                width: {{ size: 2340, type: WidthType.DXA }},
                                shading: {{ fill: "2E4053", type: ShadingType.CLEAR }},
                                children: [new Paragraph({{
                                    alignment: AlignmentType.CENTER,
                                    children: [new TextRun({{ text: "Has OHLCV", bold: true, color: "FFFFFF", size: 22 }})]
                                }})]
                            }})
                        ]
                    }}),
                    ...exchangeData.map(ex => new TableRow({{
                        children: [
                            new TableCell({{
                                borders: {{ top: tableBorder, bottom: tableBorder, left: tableBorder, right: tableBorder }},
                                width: {{ size: 2340, type: WidthType.DXA }},
                                children: [new Paragraph({{ children: [new TextRun(String(ex.name || 'N/A'))] }})]
                            }}),
                            new TableCell({{
                                borders: {{ top: tableBorder, bottom: tableBorder, left: tableBorder, right: tableBorder }},
                                width: {{ size: 2340, type: WidthType.DXA }},
                                children: [new Paragraph({{
                                    alignment: AlignmentType.CENTER,
                                    children: [new TextRun(String(ex.total_spot_pairs ?? 'N/A'))]
                                }})]
                            }}),
                            new TableCell({{
                                borders: {{ top: tableBorder, bottom: tableBorder, left: tableBorder, right: tableBorder }},
                                width: {{ size: 2340, type: WidthType.DXA }},
                                children: [new Paragraph({{
                                    alignment: AlignmentType.CENTER,
                                    children: [new TextRun(String(ex.usdt_quoted_pairs ?? 'N/A'))]
                                }})]
                            }}),
                            new TableCell({{
                                borders: {{ top: tableBorder, bottom: tableBorder, left: tableBorder, right: tableBorder }},
                                width: {{ size: 2340, type: WidthType.DXA }},
                                children: [new Paragraph({{
                                    alignment: AlignmentType.CENTER,
                                    children: [new TextRun({{ text: ex.supports_fetchOHLCV ? "Yes" : "No" }})]
                                }})]
                            }})
                        ]
                    }})))
                ]
            }}),
            
            // Technical Analysis
            new Paragraph({{
                heading: HeadingLevel.HEADING_1,
                children: [new TextRun("Technical Analysis Details")]
            }}),
            ...suggestions.map(s => new Paragraph({{
                spacing: {{ after: 150 }},
                numbering: {{ reference: "bullet-list", level: 0 }},
                children: [new TextRun(s)]
            }})),
            
            // Chart (if available)
            ...(chartPath ? [
                new Paragraph({{
                    heading: HeadingLevel.HEADING_2,
                    spacing: {{ before: 240 }},
                    children: [new TextRun("Price Chart")]
                }}),
                new Paragraph({{
                    alignment: AlignmentType.CENTER,
                    spacing: {{ before: 120, after: 240 }},
                    children: [new ImageRun({{
                        type: 'png',
                        data: fs.readFileSync(chartPath),
                        transformation: {{ width: 600, height: 400 }},
                        altText: {{ title: 'Technical Analysis Chart', description: 'Price chart with indicators', name: 'Chart' }}
                    }})]
                }})
            ] : []),
            
            // Educational Content
            new Paragraph({{
                heading: HeadingLevel.HEADING_1,
                children: [new TextRun("Understanding Technical Indicators")]
            }}),
            new Paragraph({{
                spacing: {{ after: 180 }},
                children: [new TextRun("Learn about the indicators used in this analysis:")]
            }}),
            new Paragraph({{
                heading: HeadingLevel.HEADING_2,
                children: [new TextRun("Moving Averages")]
            }}),
            new Paragraph({{
                numbering: {{ reference: "bullet-list", level: 0 }},
                children: [new TextRun("Smooth price data to identify trends")]
            }}),
            new Paragraph({{
                numbering: {{ reference: "bullet-list", level: 0 }},
                children: [new TextRun("Golden Cross: Short MA > Long MA (bullish)")]
            }}),
            new Paragraph({{
                numbering: {{ reference: "bullet-list", level: 0 }},
                children: [new TextRun("Death Cross: Short MA < Long MA (bearish)")]
            }}),
            
            new Paragraph({{
                heading: HeadingLevel.HEADING_2,
                spacing: {{ before: 200 }},
                children: [new TextRun("RSI (Relative Strength Index)")]
            }}),
            new Paragraph({{
                numbering: {{ reference: "bullet-list", level: 0 }},
                children: [new TextRun("Measures momentum (0-100 scale)")]
            }}),
            new Paragraph({{
                numbering: {{ reference: "bullet-list", level: 0 }},
                children: [new TextRun("RSI > 70: Overbought (potential pullback)")]
            }}),
            new Paragraph({{
                numbering: {{ reference: "bullet-list", level: 0 }},
                children: [new TextRun("RSI < 30: Oversold (potential bounce)")]
            }}),
            
            // Disclaimer
            new Paragraph({{
                heading: HeadingLevel.HEADING_2,
                spacing: {{ before: 240 }},
                children: [new TextRun({{ text: "Important Disclaimer", color: "E74C3C" }})]
            }}),
            new Paragraph({{
                numbering: {{ reference: "bullet-list", level: 0 }},
                children: [new TextRun("This report is for educational purposes only")]
            }}),
            new Paragraph({{
                numbering: {{ reference: "bullet-list", level: 0 }},
                children: [new TextRun("NOT financial advice")]
            }}),
            new Paragraph({{
                numbering: {{ reference: "bullet-list", level: 0 }},
                children: [new TextRun("Cryptocurrency trading has substantial risk")]
            }}),
            new Paragraph({{
                numbering: {{ reference: "bullet-list", level: 0 }},
                children: [new TextRun("Always do your own research (DYOR)")]
            }}),
            new Paragraph({{
                numbering: {{ reference: "bullet-list", level: 0 }},
                children: [new TextRun("Never invest more than you can afford to lose")]
            }})
        ]
    }}]
}});

// Save document
Packer.toBuffer(doc)
    .then(buffer => {{
        fs.writeFileSync(outputPath, buffer);
        console.log("✓ Word document generated successfully");
    }})
    .catch(err => {{
        console.error("✗ Error generating Word document:", err.message);
        process.exit(1);
    }});
"""

        return js_code